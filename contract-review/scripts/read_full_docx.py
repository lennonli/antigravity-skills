
import sys
from docx import Document

def read_docx(file_path):
    try:
        doc = Document(file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        # Also read tables as they often contain important contract info
        for table in doc.tables:
            for row in table.rows:
                text += '\n' + ' | '.join([cell.text for cell in row.cells])
        print(text)
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_contract.py <path>")
        sys.exit(1)
    read_docx(sys.argv[1])
