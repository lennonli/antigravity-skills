import sys
import os
import argparse
import time
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

# Helper to create w:ins (Insertion)
def create_ins_element(text, author="ABL-LICHENG", date_str=None):
    if date_str is None:
        date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
        
    run = OxmlElement('w:r')
    text_el = OxmlElement('w:t')
    text_el.text = text
    run.append(text_el)
    
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(int(time.time() * 1000))) # Unique ID
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date_str)
    ins.append(run)
    return ins

# Helper to create w:del (Deletion)
def create_del_element(text, author="ABL-LICHENG", date_str=None):
    if date_str is None:
        date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
        
    run = OxmlElement('w:r')
    del_text = OxmlElement('w:delText')
    del_text.text = text
    run.append(del_text)
    
    del_el = OxmlElement('w:del')
    del_el.set(qn('w:id'), str(int(time.time() * 1000) + 1))
    del_el.set(qn('w:author'), author)
    del_el.set(qn('w:date'), date_str)
    del_el.append(run)
    return del_el

def apply_revisions(doc_path, output_path, revisions):
    """
    Apply revisions to a DOCX file using native Track Changes.
    
    revisions: List of dicts {'original': 'text to find', 'revised': 'new text'}
    """
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error loading document: {e}")
        return

    # Enable tracking changes view in the saved file? 
    # doc.settings.element.append(OxmlElement('w:trackRevisions')) # Optional: force tracking on
    
    replacements_made = 0
    
    for rev in revisions:
        original_text = rev.get('original', '').strip()
        revised_text = rev.get('revised', '')
        
        if not original_text:
            continue
            
        found = False
        for p in doc.paragraphs:
            if original_text in p.text:
                # Found the paragraph containing the text
                # To properly implement track changes, we need to rebuild the paragraph components
                # This simple implementation replaces the WHOLE occurrence + surrounding text in that paragraph
                # while preserving the "track changes" structure for the part that changed.
                
                # Split paragraph text by the original_text
                # Note: This method loses character-level formatting (bold/italic) of the paragraph
                # but ensures the revision is strictly recorded.
                parts = p.text.split(original_text)
                
                # Clear existing runs
                p.clear() 
                
                # Rebuild
                for i, part in enumerate(parts):
                    # Add the unchanged text before/between matches
                    if part:
                        run = p.add_run(part)
                    
                    # Add the revision (Delete Original + Insert New)
                    if i < len(parts) - 1:
                        # 1. Add Deletion Marker
                        p._element.append(create_del_element(original_text))
                        
                        # 2. Add Insertion Marker (if there is new text)
                        if revised_text:
                            p._element.append(create_ins_element(revised_text))
                            
                found = True
                replacements_made += 1
                # Break after one replacement per revision to avoid infinite loops if original == revised
                # Remove break if you want global replace
                break 
        
        if not found:
            print(f"Warning: Could not find text: '{original_text[:30]}...'")

    try:
        doc.save(output_path)
        print(f"Successfully saved revised contract to: {output_path}")
        print(f"Total revisions applied: {replacements_made}")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Apply Track Changes revisions to a contract.")
    parser.add_argument("input_file", help="Path to original .docx file")
    parser.add_argument("--revisions", required=True, help="List of 'original'|'revised' pairs separated by ;; (e.g. 'old text'|'new text';;'foo'|'bar')")
    parser.add_argument("--output", help="Output file path (optional)")
    parser.add_argument("--open", action="store_true", help="Open the file after saving (macOS only)")
    
    args = parser.parse_args()
    
    # Parse revisions string
    # Format: "original text"|"new text";;"original text 2"|"new text 2"
    revisions_list = []
    raw_revs = args.revisions.split(';;')
    for r in raw_revs:
        if '|' in r:
            parts = r.split('|')
            # Handle cases where value might contain |
            orig = parts[0]
            new = '|'.join(parts[1:]) 
            revisions_list.append({'original': orig, 'revised': new})
            
    if not args.output:
        base, ext = os.path.splitext(args.input_file)
        args.output = f"{base}_revised{ext}"
        
    apply_revisions(args.input_file, args.output, revisions_list)

    if args.open:
        try:
            import subprocess
            print(f"Opening {args.output}...")
            subprocess.run(['open', args.output], check=True)
        except Exception as e:
            print(f"Error opening file: {e}")
