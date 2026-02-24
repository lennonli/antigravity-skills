import os
import sys
import argparse
import subprocess
import platform
from pathlib import Path
import docx
from docx.oxml.ns import qn
from pypdf import PdfReader
import google.generativeai as genai
from zhipuai import ZhipuAI
from typing import List, Optional

class DocumentTranslator:
    def __init__(self, api_type: str = "gemini", api_key: Optional[str] = None):
        self.api_type = api_type.lower()
        self.api_key = api_key or os.getenv("GOOGLE_API_KEY") if self.api_type == "gemini" else os.getenv("ZHIPU_API_KEY")
        
        if not self.api_key:
            raise ValueError(f"API key for {self.api_type} not found. Please set GOOGLE_API_KEY or ZHIPU_API_KEY environment variable.")
        
        if self.api_type == "gemini":
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel('gemini-1.5-pro')
        elif self.api_type == "zhipu":
            self.client = ZhipuAI(api_key=self.api_key)
        else:
            raise ValueError(f"Unsupported API type: {self.api_type}")

    def extract_text(self, file_path: str) -> str:
        suffix = Path(file_path).suffix.lower()
        if suffix in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif suffix == ".docx":
            doc = docx.Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        elif suffix == ".pdf":
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
            raise ValueError(f"Unsupported file format: {suffix}")

    def translate_text(self, text: str, target_lang: str = "Chinese") -> str:
        if not text.strip():
            return ""
        prompt = f"Translate the following text into {target_lang}. Maintain the original tone and formatting as much as possible. Only return the translated text.\n\nText:\n{text}"
        
        if self.api_type == "gemini":
            response = self.model.generate_content(prompt)
            return response.text.strip()
        elif self.api_type == "zhipu":
            response = self.client.chat.completions.create(
                model="glm-4",
                messages=[
                    {"role": "user", "content": prompt}
                ],
            )
            return response.choices[0].message.content.strip()
        return ""

    def open_file(self, file_path: str):
        if platform.system() == "Darwin":  # macOS
            subprocess.run(["open", file_path])
        elif platform.system() == "Windows":  # Windows
            os.startfile(file_path)
        else:  # Linux
            subprocess.run(["xdg-open", file_path])

    def translate_numbering(self, doc, target_lang: str):
        try:
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                return
                
            numbering_xml = numbering_part._element
            lvls = numbering_xml.xpath('//w:lvl')
            for lvl in lvls:
                # 1. Change number format if it's a Chinese format and we are translating to non-Chinese
                numFmt = lvl.find(qn('w:numFmt'))
                if numFmt is not None:
                    val = numFmt.get(qn('w:val'))
                    chinese_fmts = ["chineseCounting", "chineseLegalCounting", "ideographTraditional", "ideographZodiac", "ideographZodiacTraditional", "taiwaneseCounting", "taiwaneseCountingThousand", "chineseCountingThousand", "chineseCountingTenThousand"]
                    if val in chinese_fmts:
                        if "Chinese" not in target_lang:
                            numFmt.set(qn('w:val'), 'decimal')
                
                # 2. Translate prefix/suffix text
                lvlText = lvl.find(qn('w:lvlText'))
                if lvlText is not None:
                    val = lvlText.get(qn('w:val'))
                    if val:
                        new_val = val
                        if "Chinese" not in target_lang:
                            # Simple heuristics for common Chinese legal/document numbering
                            if "第" in val and "条" in val:
                                new_val = val.replace("第", "Article ").replace("条", "").strip()
                            elif "第" in val and "章" in val:
                                new_val = val.replace("第", "Chapter ").replace("章", "").strip()
                            elif "第" in val and "节" in val:
                                new_val = val.replace("第", "Section ").replace("节", "").strip()
                        
                        if new_val != val:
                            # To handle spacing nicely, e.g., 'Article %1.' we might append a dot if it lacks it and we removed Chinese chars context
                            if new_val.endswith("%1") or new_val.endswith("%2") or new_val.endswith("%3") or new_val.endswith("%4"):
                                new_val += "."
                            lvlText.set(qn('w:val'), new_val)
        except Exception as e:
            print(f"Warning: Failed to translate numbering XML: {e}")

    def process(self, input_file: str, output_file: str, target_lang: str = "Chinese", skip_open: bool = False, progress_callback=None):
        suffix = Path(input_file).suffix.lower()
        
        if suffix == ".docx":
            print(f"Modifying {input_file} in-place...")
            doc = docx.Document(input_file)
            
            # Translate auto-numberings first
            self.translate_numbering(doc, target_lang)
            
            # Count translatable paragraphs
            translatable_paras = [p for p in doc.paragraphs if p.text.strip()]
            total_paras = len(translatable_paras)
            
            print(f"Translating {total_paras} paragraphs using {self.api_type}...")
            
            processed_count = 0
            for p in doc.paragraphs:
                original_text = p.text
                if original_text.strip():
                    processed_count += 1
                    print(f"  Translating paragraph {processed_count}/{total_paras}...")
                    
                    translated_text = self.translate_text(original_text, target_lang)
                    
                    # Clear runs and add translated text to the first run to preserve paragraph style
                    # If no runs exist, just append text
                    if p.runs:
                        p.runs[0].text = translated_text
                        for i in range(1, len(p.runs)):
                            p.runs[i].text = ""
                    else:
                        p.add_run(translated_text)
                        
                    if progress_callback:
                        progress_callback(processed_count, total_paras)
            
            # Also process tables
            # This is a basic implementation. For nested tables, recursion would be needed.
            table_cells = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_cells.append(cell)
            
            total_cells = len(table_cells)
            if total_cells > 0:
                print(f"Translating {total_cells} table cells...")
                cell_count = 0
                for cell in table_cells:
                    cell_count += 1
                    print(f"  Translating cell {cell_count}/{total_cells}...")
                    
                    original_text = cell.text
                    translated_text = self.translate_text(original_text, target_lang)
                    
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = translated_text
                        for i in range(1, len(cell.paragraphs[0].runs)):
                            cell.paragraphs[0].runs[i].text = ""
                    else:
                        cell.text = translated_text
                        
                    if progress_callback:
                        # Continue progress bar after paragraphs
                        progress_callback(processed_count + cell_count, total_paras + total_cells)
                        
            print(f"Saving formatted translation to {output_file}...")
            doc.save(output_file)
            
        else:
            print(f"Extracting raw text from {input_file}...")
            text = self.extract_text(input_file)
            
            if progress_callback:
                progress_callback(0, 1) # Started
                
            print(f"Translating using {self.api_type}...")
            translated_text = self.translate_text(text, target_lang)
            
            if progress_callback:
                progress_callback(1, 1) # Done
            
            print(f"Saving translation to {output_file}...")
            doc = docx.Document()
            for p_text in translated_text.split("\n"):
                if p_text.strip():
                    doc.add_paragraph(p_text.strip())
            doc.save(output_file)
        
        if not skip_open:
            print(f"Opening {output_file}...")
            self.open_file(output_file)
        print("Done!")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Translate documents using Gemini or Zhipu AI.")
    parser.add_argument("input", help="Path to input file (txt, docx, pdf)")
    parser.add_argument("--output", help="Path to output docx file")
    parser.add_argument("--api", default="gemini", choices=["gemini", "zhipu"], help="AI API to use")
    parser.add_argument("--lang", default="Chinese", help="Target language")
    
    args = parser.parse_args()
    
    output_path = args.output or str(Path(args.input).with_suffix(".translated.docx"))
    
    try:
        translator = DocumentTranslator(api_type=args.api)
        translator.process(args.input, output_path, args.lang)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
